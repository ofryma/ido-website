import os
from doc2docx import convert
from pathlib import Path

import tempfile
from typing import Literal
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from fastapi import UploadFile
from fastapi.responses import FileResponse
from googletrans import Translator
from sqlalchemy import create_engine, Column, Integer, String
from sqlalchemy.orm import declarative_base, sessionmaker , Session


# Define the base class for SQLAlchemy models
Base = declarative_base()

SUPPORTED_LNGS = ["he"]

# Define the Translation model
class Dictionary(Base):

    __tablename__ = "dictionary"

    id = Column(Integer, primary_key=True, autoincrement=True)
    word = Column(String, nullable=False)
    translation = Column(String, nullable=False)
    
# Define the LawTranslator class
class LawTranslator:

    def __init__(self):

        self.db_engines = {}

        for lng in SUPPORTED_LNGS:
                
            db_path = f'sqlite:///{os.getcwd()}/{lng}.db'
            self.db_engines[lng] = create_engine(db_path)
            self.create_db(lng)

    def _check_supported_lng(self , lng: str):

        if not lng in SUPPORTED_LNGS:
            raise ValueError(f"Error: language {lng} is not supported. please use one of the following: {' | '.join(SUPPORTED_LNGS)}")

    def create_db(self , lng: str):
        """Creates the database and tables if they do not exist."""
        Base.metadata.create_all(self.db_engines[lng])

    def _translate(
            self, 
            text: str, 
            src_lng: Literal["he", "en", "auto"], 
            dest_lng: Literal["he", "en", "auto"]
    ) -> str:

        self._check_supported_lng(src_lng)

        try:

            lng_dict = self.fetch_dictionary(lng=src_lng)

            for key , value in lng_dict.items():
                text = text.replace(key , value)

            translator = Translator()
            translation = translator.translate(text, src=src_lng, dest=dest_lng)
            translated_text: str = translation.text.replace('.', '. ')

            return translated_text
        except: 
            return text

    def _get_session(self , src_lng: str) -> Session:

        self._check_supported_lng(src_lng)

        return sessionmaker(bind=self.db_engines[src_lng])()

    def fetch_dictionary(self , lng: str) -> dict:

        session: Session = self._get_session(src_lng=lng)

        records: list[Dictionary] = session.query(Dictionary).all()

        dictionary = {}

        for record in records:
            dictionary[record.word] = record.translation

        return dictionary

    def add_to_dictionary(self , word: str , translation: str , src_lng: str):

        session: Session = self._get_session(src_lng=src_lng)

        new_word_translation = Dictionary(
            word = word,
            translation = translation
        )

        session.add(new_word_translation)
        session.commit()
        session.close()

    def _convert_doc_to_docx(self, input_path: str) -> str:
        """
        Converts a .doc file to .docx using the doc2docx library (Windows only).
        Returns the path to the new .docx file.
        """
        input_path : Path = Path(input_path)
        if input_path.suffix.lower() != ".doc":
            raise ValueError("Input file must be .doc")

        output_path = input_path.with_suffix(".docx")

        try:

            convert(str(input_path), str(output_path))
        except Exception as e:
            raise RuntimeError("Failed to convert .doc to .docx using doc2docx") from e

        if not output_path.exists():
            raise FileNotFoundError("Converted .docx file was not created")

        return str(output_path)

    def _flip_table_direction(self, tbl: Table):
        # Flip the table layout direction to LTR
        tbl_element = tbl._tbl  # access the low-level XML element
        tblPr = tbl_element.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl_element.insert(0, tblPr)

        bidiVisual = tblPr.find(qn('w:bidiVisual'))
        if bidiVisual is not None:
            tblPr.remove(bidiVisual)

    def _translate_document_paragraph(self , para: Paragraph):
            
            # Set paragraph direction to Left-to-Right
            p = para._p
            pPr = p.get_or_add_pPr()
            bidi = pPr.find(qn('w:bidi'))
            if bidi is not None:
                pPr.remove(bidi)
            element = OxmlElement('w:rtl')
            element.set(qn('w:val'), '0')
            pPr.append(element)

            # Translate each run and preserve formatting
            for run in para.runs:
                translated_text = self.translate_hebrew_to_english(run.text)
                print(translated_text)
                run.text = translated_text

    async def translate_document(self, document: UploadFile) -> FileResponse:
        original_extension = Path(document.filename).suffix.lower()

        # Save uploaded file temporarily
        with tempfile.NamedTemporaryFile(delete=False, suffix=original_extension) as temp_file:
            temp_path = temp_file.name
            content = await document.read()
            temp_file.write(content)

        # If it's a .doc file, convert it to .docx
        if original_extension == ".doc":
            print("Converting .doc to .docx")
            temp_path = self._convert_doc_to_docx(temp_path)

        # Now open the .docx file
        doc = Document(temp_path)

        for para in doc.paragraphs:
            self._translate_document_paragraph(para)


        for tbl in doc.tables:
            
            self._flip_table_direction(tbl)

            for row in tbl.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._translate_document_paragraph(para)



        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as translated_temp_file:    
            translated_file_path = translated_temp_file.name
            doc.save(translated_file_path)

        return FileResponse(
            translated_file_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=f"{Path(document.filename).stem} translated.docx",
        )
    
    def translate_hebrew_to_english(self, text: str) -> str:

        if text.strip(" ") == "" or text is None:
            return text

        return self._translate(
            text=text,
            src_lng="he",
            dest_lng="en"
        )
